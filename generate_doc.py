"""
Generate a comprehensive Word document about the Deepfake Detector project.
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# ── Page Margins ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ── Style Helpers ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h

def add_body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
            for paragraph in row.cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()  # spacing
    return table

# ══════════════════════════════════════════════════════════════════
#                         TITLE PAGE
# ══════════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('DEEPFAKE DETECTOR')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Multimodal AI-Based Deepfake Detection System')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x44, 0x44, 0x66)

doc.add_paragraph()
doc.add_paragraph()

desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run('A comprehensive web application that uses multiple AI models\nto detect manipulated images, videos, and audio files.')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('Project Documentation')
run.font.size = Pt(14)
run.bold = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#                      TABLE OF CONTENTS (Manual)
# ══════════════════════════════════════════════════════════════════

add_heading_styled('Table of Contents', level=1)
toc_items = [
    '1. Project Overview',
    '2. Problem Statement',
    '3. System Architecture',
    '4. Technology Stack',
    '5. Algorithms and AI Models',
    '   5.1 EfficientNet-B0 (Visual Deepfake Detection)',
    '   5.2 Wav2Vec2 (Audio Deepfake Detection)',
    '   5.3 MTCNN (Face Detection)',
    '   5.4 FaceNet / InceptionResNetV1 (Face Recognition)',
    '   5.5 DBSCAN (Identity Clustering)',
    '   5.6 Non-Maximum Suppression (NMS)',
    '   5.7 C2PA / JUMBF (Content Provenance)',
    '   5.8 Weighted Score Aggregation',
    '6. Tools and Technologies',
    '   6.1 Google-Originated Technologies',
    '   6.2 Backend Tools (Python)',
    '   6.3 Frontend Tools (JavaScript/TypeScript)',
    '   6.4 AI/ML Libraries',
    '7. Scoring System',
    '8. API Reference',
    '9. Frontend Pages',
    '10. File Structure',
    '11. Security Measures',
    '12. Limitations',
    '13. Future Scope',
    '14. Conclusion',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#                    1. PROJECT OVERVIEW
# ══════════════════════════════════════════════════════════════════

add_heading_styled('1. Project Overview', level=1)

add_body(
    'The Deepfake Detector is a multimodal AI-powered platform that analyzes images, videos, '
    'and audio files for signs of synthetic manipulation. It determines whether media is authentic '
    'or fake using multiple deep learning models and a cryptographic provenance validator.'
)

add_body(
    'The system employs four independent AI analysis pipelines that work together to provide a '
    'comprehensive authenticity assessment:'
)

add_bullet(' Visual deepfake detection using EfficientNet-B0 (Convolutional Neural Network)', 'Visual Analysis:')
add_bullet(' Audio deepfake detection using Wav2Vec2 (Transformer-based model)', 'Audio Analysis:')
add_bullet(' Face detection and identity clustering using MTCNN + FaceNet + DBSCAN', 'Face Analysis:')
add_bullet(' Binary JUMBF parsing for C2PA content provenance verification', 'Provenance Check:')

add_body(
    'The results from all four pipelines are combined using a weighted scoring algorithm to produce '
    'a single authenticity score (0-100%) and a human-readable verdict (Authentic, Suspicious, '
    'Likely Fake, or Fake).'
)

# ══════════════════════════════════════════════════════════════════
#                    2. PROBLEM STATEMENT
# ══════════════════════════════════════════════════════════════════

add_heading_styled('2. Problem Statement', level=1)

add_body(
    'Deepfakes — AI-generated or AI-manipulated media — have become a significant cybersecurity '
    'threat. Advances in generative AI have made it increasingly easy to create highly realistic '
    'fake videos, images, and audio that are nearly indistinguishable from real content to the '
    'human eye and ear.'
)

add_body('The key challenges this project addresses include:')

add_bullet(' The proliferation of face-swap deepfakes used for identity fraud and misinformation', 'Visual Deepfakes:')
add_bullet(' AI-generated and cloned speech used for voice phishing and impersonation', 'Audio Deepfakes:')
add_bullet(' Lack of easily accessible tools for end-users to verify media authenticity', 'Accessibility:')
add_bullet(' Most existing detectors analyze only one modality; this project combines visual, audio, face, and provenance analysis', 'Multimodal Gap:')

add_body(
    'This project aims to provide an accessible, comprehensive, and accurate deepfake detection '
    'system that can be used by journalists, content moderators, law enforcement, and general '
    'users to verify the authenticity of digital media.'
)

# ══════════════════════════════════════════════════════════════════
#                    3. SYSTEM ARCHITECTURE
# ══════════════════════════════════════════════════════════════════

add_heading_styled('3. System Architecture', level=1)

add_body(
    'The system follows a client-server architecture with a clear separation between the '
    'frontend (user interface) and backend (AI processing).'
)

add_heading_styled('3.1 Architecture Overview', level=2)

add_table(
    ['Component', 'Technology', 'Port', 'Role'],
    [
        ['Frontend', 'React + TypeScript + Vite', '5173', 'User interface, file upload, results display'],
        ['Backend', 'Python + FastAPI', '8000', 'REST API, AI model inference, scoring'],
        ['Communication', 'HTTP REST (JSON)', '-', 'POST /upload with multipart/form-data'],
    ]
)

add_heading_styled('3.2 Request Flow', level=2)

add_body('When a user uploads a file, the system executes the following steps in sequence:')
add_bullet(' User uploads image/video/audio through the React frontend')
add_bullet(' Frontend sends POST /upload request with the file to the FastAPI backend')
add_bullet(' Backend determines the MIME type (image, video, or audio)')
add_bullet(' C2PA Validator checks for cryptographic provenance metadata')
add_bullet(' Visual model (EfficientNet-B0) analyzes frames for visual artifacts')
add_bullet(' Audio model (Wav2Vec2) analyzes audio track for synthetic speech')
add_bullet(' Face analyzer (MTCNN + FaceNet + DBSCAN) detects and clusters faces')
add_bullet(' Scoring engine computes weighted authenticity score')
add_bullet(' Verdict generator produces human-readable result')
add_bullet(' JSON response is sent back to the frontend')
add_bullet(' Results dashboard renders all analysis details')

add_heading_styled('3.3 Design Patterns', level=2)

add_body(
    'The backend uses a modular pipeline pattern. Each analysis type is encapsulated in its own '
    'class with a consistent interface:'
)

add_table(
    ['Class', 'File', 'Interface Method'],
    [
        ['VisualArtifactDetector', 'inference/visual_model.py', 'analyze_image(), analyze_video()'],
        ['AudioArtifactDetector', 'inference/audio_model.py', 'analyze_audio()'],
        ['C2PAValidator', 'inference/c2pa_validator.py', 'validate_manifest()'],
        ['FaceAnalyzer', 'inference/face_analysis.py', 'analyze_image(), analyze_video()'],
    ]
)

add_body(
    'This modular design allows each model to be independently updated, tested, or replaced '
    'without affecting the rest of the system.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#                    4. TECHNOLOGY STACK
# ══════════════════════════════════════════════════════════════════

add_heading_styled('4. Technology Stack', level=1)

add_heading_styled('4.1 Backend Technologies', level=2)

add_table(
    ['Technology', 'Version', 'Purpose'],
    [
        ['Python', '3.10+', 'Core programming language'],
        ['FastAPI', '0.110.0', 'Async web framework with auto-docs'],
        ['Uvicorn', '0.29.0', 'ASGI server for FastAPI'],
        ['PyTorch', '2.2.1', 'Deep learning framework (Meta AI)'],
        ['torchvision', '0.17.1', 'Image transforms and model architectures'],
        ['torchaudio', '2.2.1', 'Audio processing support'],
        ['HuggingFace Transformers', '≥4.38.0', 'Pre-trained model loading (Wav2Vec2)'],
        ['OpenCV (headless)', '4.9.0.80', 'Video frame extraction'],
        ['Pillow', '10.2.0', 'Image processing'],
        ['librosa', '0.10.1', 'Audio loading and resampling'],
        ['scikit-learn', '≥1.4.0', 'DBSCAN clustering'],
        ['NumPy', '1.26.4', 'Numerical computing'],
        ['Pydantic', '-', 'Request/response validation'],
    ]
)

add_heading_styled('4.2 Frontend Technologies', level=2)

add_table(
    ['Technology', 'Purpose'],
    [
        ['React + React-DOM', 'Component-based UI framework (by Meta)'],
        ['TypeScript', 'Type-safe JavaScript (by Microsoft)'],
        ['Vite', 'Fast build tool and dev server'],
        ['TailwindCSS', 'Utility-first CSS framework'],
        ['React Router DOM', 'Client-side routing'],
        ['Axios', 'HTTP client for API calls'],
        ['Lucide React', 'Icon library'],
        ['clsx + tailwind-merge', 'Conditional CSS class utilities'],
        ['Node.js', 'JavaScript runtime (V8 engine by Google)'],
        ['npm', 'Package manager'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#                    5. ALGORITHMS AND AI MODELS
# ══════════════════════════════════════════════════════════════════

add_heading_styled('5. Algorithms and AI Models', level=1)

add_body(
    'The system uses 8 distinct algorithms and models. This section provides a detailed '
    'explanation of each.'
)

add_table(
    ['#', 'Algorithm', 'Type', 'Purpose'],
    [
        ['1', 'EfficientNet-B0', 'Convolutional Neural Network', 'Visual deepfake detection'],
        ['2', 'Wav2Vec2', 'Transformer', 'Audio deepfake detection'],
        ['3', 'MTCNN (PNet → RNet → ONet)', '3-stage CNN cascade', 'Face detection'],
        ['4', 'FaceNet (InceptionResNetV1)', 'CNN', 'Face embedding (512-dim)'],
        ['5', 'DBSCAN', 'Density-based clustering', 'Identity clustering'],
        ['6', 'Non-Maximum Suppression', 'Box filtering algorithm', 'Remove duplicate detections'],
        ['7', 'C2PA / JUMBF Parser', 'Binary protocol parser', 'Content provenance verification'],
        ['8', 'Weighted Score Aggregation', 'Scoring formula', 'Combine model outputs'],
    ]
)

# ── 5.1 EfficientNet ──
add_heading_styled('5.1 EfficientNet-B0 — Visual Deepfake Detection', level=2)

add_body(
    'EfficientNet is a family of convolutional neural networks developed by Google Brain in 2019. '
    'The key innovation is compound scaling — instead of arbitrarily scaling network width, depth, '
    'or resolution independently, EfficientNet uses a mathematically optimized compound coefficient '
    'that scales all three dimensions simultaneously.'
)

add_body('Paper: "EfficientNet: Rethinking Model Scaling for Convolutional Neural Networks" '
         'by Mingxing Tan & Quoc V. Le (Google Brain, 2019)')

add_body('Key architectural features:')
add_bullet(' MBConv blocks (Mobile Inverted Bottleneck Convolutions) — efficient convolutional blocks')
add_bullet(' Depthwise separable convolutions — reduce computation by factoring standard convolutions')
add_bullet(' Squeeze-and-Excitation (SE) blocks — channel-wise attention mechanism')
add_bullet(' Inverted residual connections — skip connections for better gradient flow')
add_bullet(' Input size: 224 × 224 RGB images')
add_bullet(' Output: 2-class softmax (Real / Fake)')
add_bullet(' Parameters: ~5.3 million')

add_heading_styled('Training Data — FaceForensics++', level=3)

add_body(
    'The EfficientNet-B0 model used in this project was fine-tuned on the FaceForensics++ dataset, '
    'a benchmark for face manipulation detection containing:'
)
add_bullet(' 1,000 original video sequences from YouTube')
add_bullet(' 4,000+ manipulated videos using 4 methods:')
add_bullet('   - DeepFakes (autoencoder-based face swapping)')
add_bullet('   - Face2Face (expression transfer via 3D face reconstruction)')
add_bullet('   - FaceSwap (graphics-based face replacement)')
add_bullet('   - NeuralTextures (GAN-based facial reenactment)')
add_bullet(' C23 compression level (medium JPEG quality, realistic for consumer media)')

add_heading_styled('Image Preprocessing Pipeline', level=3)
add_bullet(' Resize to 224 × 224 pixels')
add_bullet(' Convert to tensor (pixel values 0-255 → 0.0-1.0)')
add_bullet(' Normalize with ImageNet statistics: mean=[0.485, 0.456, 0.406], std=[0.229, 0.224, 0.225]')
add_bullet(' Add batch dimension: shape becomes [1, 3, 224, 224]')

add_heading_styled('Video Analysis Method', level=3)
add_body(
    'For videos, the system uses uniform frame sampling: up to 32 frames are selected evenly '
    'across the video using numpy.linspace. Each frame is independently analyzed, and scores are '
    'aggregated using a blended formula:'
)
add_body('final_score = 0.6 × mean(frame_scores) + 0.4 × max(frame_scores)')
add_body(
    'The 40% weight on the maximum score ensures that even brief manipulated segments in a long '
    'video are not diluted by the many authentic frames.'
)

# ── 5.2 Wav2Vec2 ──
add_heading_styled('5.2 Wav2Vec2 — Audio Deepfake Detection', level=2)

add_body(
    'Wav2Vec2 is a self-supervised speech representation model developed by Meta AI (Facebook) '
    'in 2020. It learns speech features directly from raw audio waveforms without requiring '
    'transcribed labels during pre-training.'
)

add_body('Paper: "wav2vec 2.0: A Framework for Self-Supervised Learning of Speech Representations" '
         'by Alexei Baevski et al. (Meta AI, 2020)')

add_body('Architecture:')
add_bullet(' Feature Encoder: 7-layer 1D CNN that converts raw waveform to latent representations')
add_bullet(' Transformer Encoder: 12 layers with 768-dimensional hidden states for contextual modeling')
add_bullet(' Quantization Module: Discretizes latent representations for contrastive learning')
add_bullet(' Classification Head: 2-class output (Real / Fake) added during fine-tuning')

add_heading_styled('Self-Supervised Pre-training', level=3)
add_body(
    'During pre-training, Wav2Vec2 masks spans of the latent speech representations and trains '
    'the Transformer to identify the correct quantized representation for each masked position '
    'using a contrastive loss function. This forces the model to learn deep structural patterns '
    'in speech audio.'
)

add_heading_styled('Training Data — ASVspoof 2019', level=3)
add_body(
    'The model was fine-tuned on the ASVspoof (Automatic Speaker Verification Spoofing) 2019 '
    'challenge dataset, which contains:'
)
add_bullet(' Logical Access (LA) attacks: Text-to-Speech synthesis and voice conversion')
add_bullet(' Physical Access (PA) attacks: Replay attacks')
add_bullet(' 19 different spoofing attack types with varying sophistication levels')

add_heading_styled('Audio Processing Pipeline', level=3)
add_bullet(' Extract audio from any media file using librosa → mono, 16kHz sample rate')
add_bullet(' Truncate to maximum 30 seconds')
add_bullet(' Split into 5-second segments (skip segments shorter than 1 second)')
add_bullet(' Each segment processed by AutoFeatureExtractor (normalizes waveform)')
add_bullet(' Run through Wav2Vec2 → softmax probabilities')
add_bullet(' Aggregate: final_score = 0.6 × avg(segment_scores) + 0.4 × max(segment_scores)')

# ── 5.3 MTCNN ──
add_heading_styled('5.3 MTCNN — Multi-task Cascaded Convolutional Networks', level=2)

add_body(
    'MTCNN is a face detection algorithm that uses three cascaded neural networks, each '
    'progressively refining face detections with increasing accuracy.'
)

add_body('Paper: "Joint Face Detection and Alignment using Multi-task Cascaded Convolutional Networks" '
         'by Kaipeng Zhang et al. (2016)')

add_heading_styled('Stage 1 — P-Net (Proposal Network)', level=3)
add_body(
    'A small fully convolutional network that scans the image at multiple scales using an image '
    'pyramid (scale factor 0.709). Generates thousands of rough face candidate regions.'
)
add_bullet(' Architecture: Conv2d(3→10, 3×3) + PReLU → MaxPool(2×2) → Conv2d(10→16, 3×3) + PReLU → Conv2d(16→32, 3×3) + PReLU')
add_bullet(' Output Head 1: Conv2d(32→2, 1×1) → Face/Not-Face classification (softmax)')
add_bullet(' Output Head 2: Conv2d(32→4, 1×1) → Bounding box regression')
add_bullet(' Threshold: 0.7 (loose — fast but many false positives)')

add_heading_styled('Stage 2 — R-Net (Refinement Network)', level=3)
add_body(
    'Takes each P-Net candidate, resizes to 24×24, and classifies more accurately.'
)
add_bullet(' Architecture: Conv2d layers → Dense(576→128) → Dense(128→2) face/not-face + Dense(128→4) box regression')
add_bullet(' Threshold: 0.8 (medium — filters most false positives)')

add_heading_styled('Stage 3 — O-Net (Output Network)', level=3)
add_body(
    'Final refinement with facial landmark localization. Resizes candidates to 48×48.'
)
add_bullet(' Architecture: Conv2d layers → Dense(1152→256) → 3 output heads')
add_bullet(' Output Head 1: Dense(256→2) → Face/Not-Face')
add_bullet(' Output Head 2: Dense(256→4) → Bounding box regression')
add_bullet(' Output Head 3: Dense(256→10) → 5 facial landmarks (left eye, right eye, nose, left mouth, right mouth)')
add_bullet(' Threshold: 0.9 (strict — highest accuracy)')

add_heading_styled('Image Pyramid', level=3)
add_body(
    'MTCNN creates a scale pyramid to detect faces of different sizes. The image is resized '
    'to multiple scales using a factor of 0.709, and P-Net runs on each scale. The minimum face '
    'size is computed as: minsize = max(96 × min(width, height) / 1080, 40) pixels.'
)

# ── 5.4 FaceNet ──
add_heading_styled('5.4 FaceNet / InceptionResNetV1 — Face Recognition', level=2)

add_body(
    'FaceNet, developed by Google in 2015, maps face images to a compact 512-dimensional '
    'Euclidean space where distances between vectors directly correspond to face similarity.'
)

add_body('Paper: "FaceNet: A Unified Embedding for Face Recognition and Clustering" '
         'by Florian Schroff, Dmitry Kalenichenko, James Philbin (Google, 2015)')

add_body('Architecture — InceptionResNetV1:')
add_bullet(' Stem: Conv2d(3→32) → Conv2d(32→32) → Conv2d(32→64) → MaxPool → Conv2d(64→80) → Conv2d(80→192) → Conv2d(192→256)')
add_bullet(' 5× Block35 (Inception-ResNet-A): 3 parallel branches with residual connections, scale=0.17')
add_bullet(' Mixed_6a (Reduction-A): Downsamples from 256 to 896 channels')
add_bullet(' 10× Block17 (Inception-ResNet-B): 2 branches with 1×7/7×1 factorized convolutions, scale=0.10')
add_bullet(' Mixed_7a (Reduction-B): Downsamples from 896 to 1792 channels')
add_bullet(' 5× Block8 (Inception-ResNet-C): 2 branches with 1×3/3×1 factorized convolutions, scale=0.20')
add_bullet(' Final: AdaptiveAvgPool → Dropout(0.6) → Linear(1792→512) → BatchNorm → L2-Normalize')
add_bullet(' Output: 512-dimensional unit vector (face embedding)')

add_heading_styled('Preprocessing', level=3)
add_bullet(' Resize face crop to 160 × 160 pixels')
add_bullet(' Normalize: (pixel - 127.5) / 128.0')
add_bullet(' Batch processing in groups of 40 for GPU efficiency')

add_heading_styled('Triplet Loss (Training Method)', level=3)
add_body(
    'FaceNet was trained using triplet loss. During training, it takes triplets of images: '
    'an anchor face, a positive (same person), and a negative (different person). The loss '
    'function pulls the anchor closer to the positive and pushes it away from the negative '
    'in the embedding space. This creates a space where cosine distance directly corresponds '
    'to face similarity.'
)

# ── 5.5 DBSCAN ──
add_heading_styled('5.5 DBSCAN — Identity Clustering', level=2)

add_body(
    'DBSCAN (Density-Based Spatial Clustering of Applications with Noise) groups face embeddings '
    'into identity clusters without requiring the number of clusters to be specified in advance.'
)

add_body('Why DBSCAN over K-means:')
add_bullet(' No need to specify K — the number of identities is unknown')
add_bullet(' Handles noise — faces that do not belong to any identity are labeled as noise (label = -1)')
add_bullet(' Works naturally with cosine distance in the 512-dimensional embedding space')

add_body('Parameters used:')
add_table(
    ['Parameter', 'Value', 'Meaning'],
    [
        ['eps', '0.35', 'Maximum cosine distance for two faces to be in the same cluster'],
        ['metric', 'cosine', 'Distance metric in 512-dim embedding space'],
        ['min_samples', '5', 'Minimum face detections to form a valid identity cluster'],
    ]
)

add_heading_styled('Complete Face Analysis Pipeline', level=3)
add_bullet(' Sample 1 frame per 30 from the video')
add_bullet(' Detect all faces with MTCNN (confidence threshold > 0.98 for video, > 0.90 for images)')
add_bullet(' Scale bounding box by 1.2× and extract face crops')
add_bullet(' Generate 512-dimensional embeddings using FaceNet')
add_bullet(' If ≥ 5 faces → run DBSCAN → group by identity')
add_bullet(' If < 5 faces → treat all as one identity')
add_bullet(' Return identity clusters with face counts, confidence scores, and base64 thumbnails')

# ── 5.6 NMS ──
add_heading_styled('5.6 Non-Maximum Suppression (NMS)', level=2)

add_body(
    'NMS is a post-processing algorithm used within each MTCNN stage to remove redundant '
    'overlapping bounding boxes. When multiple candidate boxes cover the same face, NMS keeps '
    'only the box with the highest confidence score and removes all other boxes whose '
    'Intersection over Union (IoU) exceeds a threshold.'
)

add_body('Two NMS strategies are used:')
add_bullet(' Standard IoU-based NMS (Stages 1 & 2): overlap = intersection / union')
add_bullet(' "Min" NMS (Stage 3): overlap = intersection / min(area1, area2) — better for overlapping faces')
add_body(
    'The project uses torchvision\'s batched_nms for GPU-accelerated NMS in early stages, and '
    'a custom NumPy implementation for the "Min" strategy in the final stage.'
)

# ── 5.7 C2PA ──
add_heading_styled('5.7 C2PA / JUMBF — Content Provenance Verification', level=2)

add_body(
    'C2PA (Coalition for Content Provenance and Authenticity) is an open technical standard '
    'backed by Adobe, Microsoft, Google, Intel, and major camera manufacturers. When a supported '
    'camera or software creates or edits media, it can embed a cryptographic manifest inside the '
    'file containing provenance information.'
)

add_body(
    'JUMBF (JPEG Universal Metadata Box Format, ISO 19566-5) is the container format used to '
    'embed C2PA metadata inside media files.'
)

add_heading_styled('Format-Specific Parsing', level=3)

add_table(
    ['Format', 'Location', 'Method'],
    [
        ['JPEG', 'APP11 marker segment (0xFFEB)', 'Walk through JPEG marker segments, find APP11, check for JUMBF/C2PA magic bytes'],
        ['PNG', 'caBX chunk', 'Walk through PNG chunks, find caBX, check for JUMBF content'],
        ['Video/Other', 'Generic binary scan', 'Read up to 10MB, search for known C2PA byte patterns'],
    ]
)

add_body('C2PA byte patterns scanned:')
add_bullet(' c2pa — content type identifier')
add_bullet(' jumb — JUMBF box type marker')
add_bullet(' jumd — JUMBF description type marker')
add_bullet(' c2pa.claim — C2PA claim structure')
add_bullet(' c2pa.assertions — C2PA assertions block')
add_bullet(' c2pa.signature — Cryptographic signature')

add_body('Known issuer detection: Adobe, Qualcomm, Samsung, Google, Apple, Microsoft, Truepic, Leica, Nikon, Canon, Sony')

# ── 5.8 Scoring ──
add_heading_styled('5.8 Weighted Score Aggregation', level=2)

add_body('The final authenticity score (0-100%) is computed in three steps:')

add_body('Step 1 — Weighted fake probability:')
add_body('    fake_probability = visual_score × visual_weight + audio_score × audio_weight')

add_body('Step 2 — Convert to authenticity:')
add_body('    authenticity_score = (1 - fake_probability) × 100')

add_body('Step 3 — C2PA bonus:')
add_body('    If C2PA valid signature AND hardware attestation found → add 15 points (capped at 100)')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#                    6. TOOLS AND TECHNOLOGIES
# ══════════════════════════════════════════════════════════════════

add_heading_styled('6. Tools and Technologies', level=1)

add_heading_styled('6.1 Google-Originated Technologies', level=2)

add_body(
    'Several core technologies in this project were developed or invented by Google:'
)

add_table(
    ['Technology', 'Google Origin', 'Role in Project'],
    [
        ['EfficientNet', 'Google Brain (Mingxing Tan & Quoc V. Le, 2019)', 'Visual deepfake detection CNN architecture'],
        ['FaceNet', 'Google (Florian Schroff et al., 2015)', 'Face embedding algorithm using triplet loss'],
        ['Inception Architecture', 'Google (Szegedy et al., 2014-2017)', 'InceptionResNetV1 neural network used inside FaceNet'],
        ['C2PA Standard', 'Google is a founding/steering member', 'Content provenance standard; Google Pixel phones embed C2PA'],
        ['Node.js V8 Engine', 'Google Chrome team', 'JavaScript runtime engine powering the frontend dev server'],
    ]
)

add_heading_styled('6.2 Backend Tools', level=2)

add_table(
    ['Tool', 'Creator', 'Purpose'],
    [
        ['Python 3.10+', 'Python Software Foundation', 'Core backend programming language'],
        ['FastAPI', 'Sebastián Ramírez', 'Async REST API framework with auto-documentation'],
        ['Uvicorn', 'Encode (Tom Christie)', 'ASGI server for running FastAPI'],
        ['PyTorch', 'Meta AI (Facebook)', 'Deep learning tensor computation and model inference'],
        ['torchvision', 'Meta AI', 'Image transforms, pre-trained model architectures, NMS'],
        ['torchaudio', 'Meta AI', 'Audio processing support'],
        ['HuggingFace Transformers', 'Hugging Face', 'Pre-trained model hub (loads Wav2Vec2)'],
        ['OpenCV', 'Intel (open source)', 'Video frame extraction, color conversion'],
        ['Pillow', 'Alex Clark', 'Image loading, resizing, format conversion'],
        ['librosa', 'Brian McFee', 'Audio loading from any format, resampling to 16kHz'],
        ['scikit-learn', 'INRIA (France)', 'DBSCAN clustering algorithm implementation'],
        ['NumPy', 'Community', 'Numerical array operations, frame sampling'],
        ['Pydantic', 'Samuel Colvin', 'API request/response data validation'],
    ]
)

add_heading_styled('6.3 Frontend Tools', level=2)

add_table(
    ['Tool', 'Creator', 'Purpose'],
    [
        ['React + React-DOM', 'Meta AI (Facebook)', 'Component-based UI framework, virtual DOM'],
        ['TypeScript', 'Microsoft', 'Type-safe JavaScript superset'],
        ['Vite', 'Evan You', 'Fast build tool with Hot Module Replacement'],
        ['TailwindCSS', 'Tailwind Labs', 'Utility-first CSS framework'],
        ['React Router DOM', 'Remix team', 'Client-side page routing'],
        ['Axios', 'Matt Zabriskie', 'HTTP client library for API calls'],
        ['Lucide React', 'Community', 'Icon library (Upload, Camera, Dashboard icons)'],
        ['clsx + tailwind-merge', 'Community', 'Conditional CSS class composition utilities'],
        ['Node.js', 'OpenJS Foundation (V8 by Google)', 'JavaScript runtime for dev tools'],
        ['npm', 'GitHub (Microsoft)', 'Package manager for frontend dependencies'],
    ]
)

add_heading_styled('6.4 Development Tools', level=2)

add_table(
    ['Tool', 'Purpose'],
    [
        ['Git', 'Version control'],
        ['VS Code', 'Code editor / IDE'],
        ['pip', 'Python package manager'],
        ['venv', 'Python virtual environment isolation'],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#                    7. SCORING SYSTEM
# ══════════════════════════════════════════════════════════════════

add_heading_styled('7. Scoring System', level=1)

add_heading_styled('7.1 Weights by Media Type', level=2)

add_table(
    ['Media Type', 'Visual Weight', 'Audio Weight', 'Rationale'],
    [
        ['Video', '0.55', '0.45', 'Both modalities matter; visual slightly more detectable'],
        ['Image', '1.00', '0.00', 'No audio track to analyze'],
        ['Audio', '0.00', '1.00', 'No visual frames to analyze'],
    ]
)

add_heading_styled('7.2 Verdict Thresholds', level=2)

add_table(
    ['Score Range', 'Verdict', 'Description'],
    [
        ['85 – 100', 'AUTHENTIC', 'High confidence the media is genuine'],
        ['60 – 84', 'SUSPICIOUS', 'Some indicators of manipulation detected'],
        ['35 – 59', 'LIKELY FAKE', 'Strong indicators of manipulation'],
        ['0 – 34', 'FAKE', 'Very high confidence of manipulation'],
    ]
)

add_heading_styled('7.3 Example Calculation', level=2)

add_body('Scenario: Video with visual_score = 0.12, audio_score = 0.08, no C2PA metadata')
add_body('fake_probability = 0.12 × 0.55 + 0.08 × 0.45 = 0.066 + 0.036 = 0.102')
add_body('authenticity_score = (1 - 0.102) × 100 = 89.8%')
add_body('Verdict: AUTHENTIC (89.8 ≥ 85)')

# ══════════════════════════════════════════════════════════════════
#                    8. API REFERENCE
# ══════════════════════════════════════════════════════════════════

add_heading_styled('8. API Reference', level=1)

add_heading_styled('8.1 GET / — Health Check', level=2)
add_body('Returns API status and loaded models. Used to verify the server is running.')

add_heading_styled('8.2 POST /upload — Media Analysis', level=2)
add_body('Accepts a media file via multipart/form-data and runs all AI analysis pipelines.')
add_body('Request: multipart/form-data with a "file" field')
add_body('Response: JSON object containing:')
add_bullet(' authenticity_score (float 0-100)')
add_bullet(' c2pa_valid (boolean)')
add_bullet(' visual_artifacts_detected (boolean)')
add_bullet(' audio_artifacts_detected (boolean)')
add_bullet(' face_detected (boolean)')
add_bullet(' faces_count (integer)')
add_bullet(' verdict (string)')
add_bullet(' details (object with full analysis from each model)')

add_body('Interactive API documentation is available at http://localhost:8000/docs (Swagger UI)')

# ══════════════════════════════════════════════════════════════════
#                    9. FRONTEND PAGES
# ══════════════════════════════════════════════════════════════════

add_heading_styled('9. Frontend Pages', level=1)

add_table(
    ['Page', 'Route', 'File', 'Description'],
    [
        ['Upload', '/', 'UploadPage.tsx', 'Drag-and-drop media upload supporting JPEG, PNG, MP4, MOV, WAV, MP3'],
        ['Camera', '/camera', 'CameraPage.tsx', 'Live webcam capture for instant photo analysis'],
        ['Results', '/results', 'ResultsPage.tsx', 'Analysis dashboard with scores, face clusters, C2PA status, verdict'],
    ]
)

add_body(
    'The frontend uses a dark theme with glassmorphism effects and gradient accents. '
    'Navigation is handled by a fixed top navigation bar with icons from Lucide React.'
)

# ══════════════════════════════════════════════════════════════════
#                    10. FILE STRUCTURE
# ══════════════════════════════════════════════════════════════════

add_heading_styled('10. File Structure', level=1)

structure = [
    ('deepfake-detector/', 'Root project directory'),
    ('├── deepfake-backend/', 'Python FastAPI Server'),
    ('│   ├── main.py', 'API entry point, orchestrates all models'),
    ('│   ├── inference/', 'AI inference modules'),
    ('│   │   ├── visual_model.py', 'EfficientNet-B0 inference'),
    ('│   │   ├── audio_model.py', 'Wav2Vec2 inference'),
    ('│   │   ├── c2pa_validator.py', 'Binary C2PA/JUMBF parser'),
    ('│   │   └── face_analysis.py', 'MTCNN + FaceNet + DBSCAN pipeline'),
    ('│   ├── dface/', 'dFace library (face detection & recognition)'),
    ('│   │   ├── __init__.py', 'Exports MTCNN, FaceNet'),
    ('│   │   ├── mtcnn.py', 'PNet, RNet, ONet + NMS'),
    ('│   │   └── facenet.py', 'InceptionResNetV1 + embedding'),
    ('│   ├── models/', 'Pre-trained model weights'),
    ('│   │   ├── efficientnet_b0_ffpp_c23.pth', 'Visual model (~20MB)'),
    ('│   │   ├── wav2vec2-deepfake/', 'Audio model (~360MB)'),
    ('│   │   ├── mtcnn.pt', 'MTCNN face detection weights'),
    ('│   │   └── facenet.pt', 'FaceNet embedding weights'),
    ('│   └── requirements.txt', 'Python dependencies'),
    ('├── deepfake-web-app/', 'React + Vite Frontend'),
    ('│   ├── src/', 'Source code'),
    ('│   │   ├── App.tsx', 'Router & navigation'),
    ('│   │   ├── pages/', 'Page components'),
    ('│   │   │   ├── UploadPage.tsx', 'Drag & drop upload'),
    ('│   │   │   ├── CameraPage.tsx', 'Live webcam capture'),
    ('│   │   │   └── ResultsPage.tsx', 'Analysis dashboard'),
    ('│   │   ├── index.css', 'TailwindCSS + custom styles'),
    ('│   │   └── main.tsx', 'React entry point'),
    ('│   ├── tailwind.config.js', 'Tailwind configuration'),
    ('│   └── vite.config.js', 'Vite build configuration'),
]

for path, desc in structure:
    p = doc.add_paragraph()
    run = p.add_run(path + '  ')
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run = p.add_run('— ' + desc)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    p.paragraph_format.space_after = Pt(1)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
#                    11. SECURITY MEASURES
# ══════════════════════════════════════════════════════════════════

add_heading_styled('11. Security Measures', level=1)

add_bullet(' MIME type validation — only image, video, and audio types are processed', 'File Type Checking: ')
add_bullet(' Files are saved to a temporary uploads/ directory and immediately deleted after analysis', 'Temporary Storage: ')
add_bullet(' Every model\'s analyze method is wrapped in try/except; one model failure doesn\'t crash the system', 'Error Isolation: ')
add_bullet(' FastAPI CORSMiddleware controls which origins can access the API', 'CORS Protection: ')
add_bullet(' All AI inference runs locally — no data is sent to external APIs', 'Local Processing: ')
add_bullet(' Pydantic models validate API response structure before sending to client', 'Input Validation: ')

# ══════════════════════════════════════════════════════════════════
#                    12. LIMITATIONS
# ══════════════════════════════════════════════════════════════════

add_heading_styled('12. Limitations', level=1)

add_bullet(' Only supports uploaded files, not live video streams')
add_bullet(' Without a GPU, analysis of long videos can take 30+ seconds')
add_bullet(' New deepfake techniques not present in the training datasets may not be detected')
add_bullet(' Most existing media does not yet have C2PA metadata')
add_bullet(' Results are not stored persistently (no database)')
add_bullet(' No authentication or multi-user support')
add_bullet(' Limited to specific file formats (JPEG, PNG, MP4, MOV, WAV, MP3)')

# ══════════════════════════════════════════════════════════════════
#                    13. FUTURE SCOPE
# ══════════════════════════════════════════════════════════════════

add_heading_styled('13. Future Scope', level=1)

add_bullet(' Real-time video analysis via WebSocket streaming')
add_bullet(' GPU acceleration for faster inference (CUDA support is already built in)')
add_bullet(' Model retraining on newer, more diverse deepfake datasets')
add_bullet(' Browser extension for automatic media verification while browsing')
add_bullet(' Batch processing for multiple files simultaneously')
add_bullet(' Database integration for storing analysis history')
add_bullet(' Mobile application using React Native')
add_bullet(' LLM integration for semantic analysis and natural language explanations')
add_bullet(' Explainability heatmaps showing which image regions triggered detection')
add_bullet(' Frequency domain analysis (DCT/FFT) to detect GAN-specific artifacts')
add_bullet(' Model ensemble approach with multiple visual classifiers for higher accuracy')

# ══════════════════════════════════════════════════════════════════
#                    14. CONCLUSION
# ══════════════════════════════════════════════════════════════════

add_heading_styled('14. Conclusion', level=1)

add_body(
    'The Deepfake Detector project demonstrates a comprehensive approach to media authenticity '
    'verification by combining multiple AI models and analysis techniques into a single, '
    'user-friendly web application. By leveraging state-of-the-art deep learning models '
    '(EfficientNet-B0, Wav2Vec2, MTCNN, FaceNet) along with cryptographic provenance validation '
    '(C2PA), the system provides robust multimodal detection capabilities.'
)

add_body(
    'The weighted scoring algorithm ensures that results are calibrated to the specific media '
    'type being analyzed, while the modular architecture allows individual components to be '
    'updated independently as deepfake technology evolves. The combination of visual artifact '
    'detection, audio synthesis analysis, face identity clustering, and content provenance '
    'checking creates a defense-in-depth approach that is more reliable than any single-model '
    'solution.'
)

add_body(
    'This project serves as a practical tool for combating the growing threat of synthetic media '
    'and contributes to the broader effort of maintaining trust in digital content.'
)

# ── Save ──
output_path = os.path.join(os.path.dirname(__file__), 'Deepfake_Detector_Project_Documentation.docx')
doc.save(output_path)
print(f"\n[OK] Word document saved to: {output_path}")
print(f"     File size: {os.path.getsize(output_path) / 1024:.1f} KB")
